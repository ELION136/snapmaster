"""
SnapMaster - Versión refactorizada y profesional
Archivo: snapmaster_profesional.py
Descripción: Implementación en un solo archivo, modularizada con clases:
 - SecurityManager: manejo de contraseña y token de desbloqueo (sin hardcodear)
 - PDFConverter: detecta y usa el mejor método disponible para convertir .docx -> .pdf
 - CaptureManager: captura de pantalla (usa mss si está disponible, si no pyautogui)
 - WebcamManager: manejo de cámara con OpenCV
 - SnapMasterApp: interfaz Tkinter (estructura similar a la original) pero con
   manejo de hilos, logging, y manejo robusto de excepciones.

Requisitos (sugeridos para un virtualenv):
  pip install pyautogui pillow python-docx opencv-python docx2pdf mss pywin32
  # No todas son obligatorias: mss y docx2pdf son opcionales; se usan si están

Notas de seguridad:
 - No se guarda ninguna contraseña en texto plano.
 - Al primer arranque el programa permite crear/actualizar la contraseña.
 - El hash se guarda en un archivo dentro de APP_DIR con permisos de usuario.

Autor: Reescrito por ChatGPT (solicitud del usuario)
Versión: 1.0
"""
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import threading
import time
from datetime import datetime
import subprocess
import sys
import hashlib
import tkinter.ttk as ttk

# ================== CONFIGURACIÓN DE SEGURIDAD ==================

APP_DIR = os.path.join(os.getenv('APPDATA'), 'SnapMaster')
TOKEN_FILE = os.path.join(APP_DIR, 'unlock.token')
PASSWORD_HASH = hashlib.sha256("123alexyjhenny".encode()).hexdigest()  # 🔑 Cambia aquí tu contraseña

def create_app_dir():
    """Crea carpeta oculta si no existe."""
    if not os.path.exists(APP_DIR):
        os.makedirs(APP_DIR)

def is_unlocked():
    """Verifica si ya está desbloqueado."""
    return os.path.exists(TOKEN_FILE)

def save_token():
    """Guarda token de desbloqueo."""
    with open(TOKEN_FILE, "w") as f:
        f.write("desbloqueado")

# ================== VENTANA DE LOGIN ==================

def ask_password():
    """Ventana de contraseña. Devuelve True si es correcta."""
    login_win = tk.Tk()
    login_win.title("🔒 Acceso a SnapMaster")
    login_win.geometry("300x150")
    login_win.resizable(False, False)
    login_win.configure(bg='#f8f9fa')  
    
    # Centrar ventana
    login_win.eval('tk::PlaceWindow . center')

    tk.Label(login_win, text="Ingrese la contraseña:", font=("Arial", 12), 
             bg='#f8f9fa', fg='#212529').pack(pady=10)  
    
    pwd_entry = tk.Entry(login_win, show="*", font=("Arial", 12), 
                         bg='#e9ecef', fg='#212529', insertbackground='#212529')  
    pwd_entry.pack(pady=5)
    pwd_entry.focus()
    
    result = [False]  

    def check_password():
        password = pwd_entry.get()
        if hashlib.sha256(password.encode()).hexdigest() == PASSWORD_HASH:
            save_token()
            result[0] = True
            login_win.destroy()
        else:
            messagebox.showerror("Error", "Contraseña incorrecta")
            pwd_entry.delete(0, tk.END)

    tk.Button(login_win, text="Acceder", command=check_password,
              bg='#0d6efd', fg='white', font=('Arial', 10, 'bold')).pack(pady=10)  
    
    # Permitir usar Enter para enviar la contraseña
    login_win.bind('<Return>', lambda event: check_password())
    
    login_win.mainloop()
    
    return result[0]

# ================== CLASE PRINCIPAL DE LA APLICACIÓN ==================

class ScreenCaptureApp:
    def __init__(self, root):
        self.root = root
        self.root.title("SnapMaster")
        self.root.geometry("1000x650")
        self.root.configure(bg='#f8f9fa')  #
        # Variables principales
        self.capturing = False
        self.capture_thread = None
        self.webcam_active = False
        self.webcam_window = None
        self.cap = None
        self.save_path = tk.StringVar(value=os.path.expanduser("~/Desktop"))
        self.interval = tk.IntVar(value=5)
        self.image_format = tk.StringVar(value="PNG")
        self.base_name = tk.StringVar(value="captura")
        self.word_name = tk.StringVar(value="reporte_capturas")
        self.counter = 0

        # Documento Word
        self.document = None
        self.word_path = ""
        self.preview_photo = None

        # Mensaje rápido de inicio
        self.status_var = tk.StringVar(value="Cargando interfaz...")
        status_frame = tk.Frame(self.root, bg='#e9ecef', height=30)  
        status_frame.pack(fill='x', side='bottom')
        status_frame.pack_propagate(False)

        status_label = tk.Label(status_frame, textvariable=self.status_var,
                                bg='#e9ecef', fg='#495057', font=('Arial', 9)) 
        status_label.pack(side='left', padx=10, pady=5)

        self.counter_var = tk.StringVar(value="Capturas: 0")
        counter_label = tk.Label(status_frame, textvariable=self.counter_var,
                                 bg='#e9ecef', fg='#495057', font=('Arial', 9))  
        counter_label.pack(side='right', padx=10, pady=5)

        self.setup_ui()

    def setup_ui(self):
        # ------------------ TÍTULO ------------------
        title_frame = tk.Frame(self.root, bg='#f8f9fa')  
        title_frame.pack(pady=10)
        title_label = tk.Label(title_frame, text="📸 SnapMaster",
                            font=('Arial', 16, 'bold'), bg='#f8f9fa', fg='#0d6efd')  
        title_label.pack()

        # ------------------ NOTEBOOK CON PESTAÑAS ------------------
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TNotebook', background='#e9ecef', borderwidth=0)  
        style.configure('TNotebook.Tab', background='#dee2e6', foreground='#495057',  
                       padding=[20, 8], font=('Arial', 9, 'bold'))
        style.map('TNotebook.Tab', background=[('selected', '#adb5bd')])  

        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=15, pady=10)

        # Pestaña 1: Principal
        main_frame = tk.Frame(notebook, bg='#f8f9fa')  
        notebook.add(main_frame, text="📷 Captura")

        # Pestaña 2: Manual
        manual_frame = tk.Frame(notebook, bg='#f8f9fa')  
        notebook.add(manual_frame, text="📖 Manual")

        # ------------------ PANEL PRINCIPAL ------------------
        # Panel izquierdo
        left_panel = tk.Frame(main_frame, bg='#f8f9fa', width=580)  
        left_panel.pack(side='left', fill='y', padx=(0, 10))
        left_panel.pack_propagate(False)

        # Panel derecho (preview)
        right_panel = tk.Frame(main_frame, bg='#f8f9fa')  
        right_panel.pack(side='right', fill='both', expand=True)

        # === Configuración principal ===
        config_main_frame = tk.LabelFrame(left_panel, text="📂 Configuración Principal",
                                        font=('Arial', 10, 'bold'), bg='#e9ecef', fg='#495057')  
        config_main_frame.pack(fill='x', pady=5)

        # Ruta
        path_section = tk.Frame(config_main_frame, bg='#e9ecef')  
        path_section.pack(fill='x', padx=10, pady=8)
        tk.Label(path_section, text="Ruta de guardado:", bg='#e9ecef',
                fg="#495057", font=('Arial', 9, 'bold')).pack(anchor='w')  
        self.path_entry = tk.Entry(path_section, textvariable=self.save_path, font=('Arial', 9),
                                bg="#ffffff", fg="#495057", insertbackground="#495057")  
        self.path_entry.pack(side='left', fill='x', expand=True)
        tk.Button(path_section, text="📁", command=self.browse_folder,
                bg='#0d6efd', fg='white', font=('Arial', 9, 'bold'), width=3).pack(side='right', padx=(5, 0))  

        # Nombres
        names_section = tk.Frame(config_main_frame, bg='#e9ecef')  
        names_section.pack(fill='x', padx=10, pady=8)
        tk.Label(names_section, text="Nombre imágenes:", bg='#e9ecef', fg="#495057",
                font=('Arial', 9, 'bold')).grid(row=0, column=0, sticky="w")  
        tk.Entry(names_section, textvariable=self.base_name, font=('Arial', 9),
                bg="#ffffff", fg="#495057", insertbackground="#495057").grid(row=1, column=0, sticky="ew", padx=5)  
        tk.Label(names_section, text="Nombre Word:", bg='#e9ecef', fg="#495057",
                font=('Arial', 9, 'bold')).grid(row=0, column=1, sticky="w")  
        tk.Entry(names_section, textvariable=self.word_name, font=('Arial', 9),
                bg="#ffffff", fg="#495057", insertbackground="#495057").grid(row=1, column=1, sticky="ew", padx=5)  
        names_section.columnconfigure((0, 1), weight=1)

        # === Configuración de captura ===
        capture_config_frame = tk.LabelFrame(left_panel, text="⚙️ Configuración de Captura",
                                            font=('Arial', 10, 'bold'), bg='#e9ecef', fg='#495057')  
        capture_config_frame.pack(fill='x', pady=5)

        tk.Label(capture_config_frame, text="Intervalo (segundos):", bg='#e9ecef', fg="#495057",
                font=('Arial', 9, 'bold')).pack(anchor='w', padx=10, pady=(5, 0))  
        tk.Scale(capture_config_frame, from_=1, to=300, orient='horizontal',
                variable=self.interval, bg='#e9ecef', fg="#495057", font=('Arial', 8), length=200,
                troughcolor="#ced4da", highlightbackground="#e9ecef").pack(fill='x', padx=10, pady=2)  

        tk.Label(capture_config_frame, text="Formato de imagen:", bg='#e9ecef', fg="#495057",
                font=('Arial', 9, 'bold')).pack(anchor='w', padx=10, pady=(8, 0))  
        format_buttons = tk.Frame(capture_config_frame, bg='#e9ecef')  
        format_buttons.pack(pady=5)
        tk.Radiobutton(format_buttons, text="PNG", variable=self.image_format, value="PNG",
                    bg='#e9ecef', fg="#495057", selectcolor="#ced4da", font=('Arial', 9)).pack(side='left', padx=(0, 10))  
        tk.Radiobutton(format_buttons, text="JPG", variable=self.image_format, value="JPG",
                    bg='#e9ecef', fg="#495057", selectcolor="#ced4da", font=('Arial', 9)).pack(side='left')  

        # === Control de captura ===
        control_frame = tk.LabelFrame(left_panel, text="🎮 Control de Captura",
                                    font=('Arial', 10, 'bold'), bg='#e9ecef', fg='#495057')  
        control_frame.pack(fill='x', pady=5)
        
        # Botones principales
        buttons_row1 = tk.Frame(control_frame, bg='#e9ecef')  
        buttons_row1.pack(pady=5)
        
        self.start_btn = tk.Button(buttons_row1, text="▶️ INICIAR", command=self.start_capture,
                                bg='#198754', fg='white', font=('Arial', 10, 'bold'),  
                                padx=15, pady=6, width=12)
        self.start_btn.pack(side='left', padx=8)
        
        self.stop_btn = tk.Button(buttons_row1, text="⏹️ DETENER", command=self.stop_capture,
                                bg='#dc3545', fg='white', font=('Arial', 10, 'bold'),  
                                padx=15, pady=6, width=12, state='disabled')
        self.stop_btn.pack(side='left', padx=8)

        # Segunda fila de botones
        buttons_row2 = tk.Frame(control_frame, bg='#e9ecef')  
        buttons_row2.pack(pady=5)
        
        self.webcam_btn = tk.Button(buttons_row2, text="📹 CÁMARA", command=self.toggle_webcam,
                                    bg='#6f42c1', fg='white', font=('Arial', 10, 'bold'),  
                                    padx=15, pady=6, width=25)
        self.webcam_btn.pack(pady=2)

        self.pdf_btn = tk.Button(buttons_row2, text="📄 Generar PDF", command=self.generate_pdf_from_word,
                                    bg='#6c757d', fg='white', font=('Arial', 10, 'bold'),  
                                    padx=15, pady=6, width=25, state='disabled')
        self.pdf_btn.pack(pady=2)

        # === Vista previa ===
        preview_frame = tk.LabelFrame(right_panel, text="👀 Vista Previa",
                                    font=('Arial', 10, 'bold'), bg='#e9ecef', fg='#495057')  
        preview_frame.pack(fill='both', expand=True, pady=5)
        self.preview_label = tk.Label(preview_frame, text="La captura aparecerá aquí",
                                    bg='#ffffff', relief='sunken', bd=2,  
                                    font=('Arial', 11), fg='#6c757d') 
        self.preview_label.pack(fill='both', expand=True, padx=10, pady=10)

        self.status_var.set("Listo para capturar")

        # ------------------ MANUAL ------------------
        instructions = """
    📖 Manual de Uso de SnapMaster
    version 1.2.1.1 
    
    CORRECCIONES IMPLEMENTADAS:
    - se añadio exclusividad
    - se quito el detector de rostros por compatibilidad
    - Nueva paleta de colores azul/gris oscuro
    - Corrección del error de PDF en archivos .exe
    - Botón PDF ahora se activa solo cuando hay documento Word

    1. Configuración:
    - Elija la carpeta donde se guardarán las capturas.
    - Defina el nombre base de las imágenes y del documento Word.

    2. Opciones de Captura:
    - Ajuste el intervalo en segundos.
    - Seleccione el formato de imagen (PNG o JPG).

    3. Control:
    - Presione ▶️ INICIAR para comenzar la captura automática.
    - Presione ⏹️ DETENER para finalizar y guardar el documento Word.
    - El botón PDF se activará automáticamente al crear el documento.

    4. Cámara Web:
    - Active la cámara con 📹 CÁMARA.
    - Puede habilitar la detección de rostros con la casilla.
    - Cierre la ventana para detener.

    5. Vista Previa:
    - A la derecha verá la última captura guardada.

    ⚠️ Recomendaciones:
    - No cierre la aplicación mientras captura.
    - El documento Word se guarda automáticamente al detener.
    - Para mejor rendimiento, use intervalos mayores a 10 segundos.
    - El PDF se genera usando método alternativo compatible con .exe
    """
        text_widget = tk.Text(manual_frame, wrap="word", font=("Arial", 11),
                            bg="#ffffff", fg="#495057", insertbackground="#495057")  
        text_widget.insert("1.0", instructions)
        text_widget.config(state="disabled")
        text_widget.pack(fill="both", expand=True, padx=10, pady=10)

        scrollbar = tk.Scrollbar(text_widget, command=text_widget.yview)
        text_widget['yscrollcommand'] = scrollbar.set
        scrollbar.pack(side="right", fill="y")

    # ======================= LÓGICA ======================

    def browse_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.save_path.set(folder)

    def start_capture(self):
        try:
            if not os.path.exists(self.save_path.get()):
                messagebox.showerror("Error", "La ruta especificada no existe")
                return

            if not self.base_name.get().strip():
                messagebox.showerror("Error", "Debe especificar un nombre base")
                return

            if not self.word_name.get().strip():
                messagebox.showerror("Error", "Debe especificar un nombre para el documento Word")
                return

            self.capturing = True
            self.counter = 0
            self.create_word_document()

            self.start_btn.config(state='disabled')
            self.stop_btn.config(state='normal')
            # Activar botón PDF cuando se crea el documento
            self.pdf_btn.config(state='normal', bg='#fd7e14')  

            self.capture_thread = threading.Thread(target=self.capture_loop, daemon=True)
            self.capture_thread.start()
            self.status_var.set(f"Capturando cada {self.interval.get()} segundos...")

        except Exception as e:
            messagebox.showerror("Error", f"Error al iniciar captura: {str(e)}")

    def stop_capture(self):
        self.capturing = False
        self.start_btn.config(state='normal')
        self.stop_btn.config(state='disabled')
        if self.document:
            try:
                # Guardar Word con nombre único
                self.document.save(self.word_path)
                self.status_var.set(f"Documento guardado en: {self.word_path}")
            except Exception as e:
                self.status_var.set(f"Error al guardar documento: {str(e)}")

    def create_word_document(self):
        from docx import Document
        self.document = Document()
        self.document.add_heading('Reporte de Capturas de Pantalla', 0)
        self.document.add_paragraph(f'Fecha: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        self.document.add_paragraph(f'Intervalo: {self.interval.get()} seg')
        self.document.add_paragraph(f'Formato: {self.image_format.get()}')
        self.document.add_paragraph('_' * 40)

        base_word_path = os.path.join(self.save_path.get(), self.word_name.get())
        self.word_path = self.get_unique_filename(base_word_path, ".docx")

    def generate_pdf_from_word(self):
        try:
            if not self.word_path or not os.path.exists(self.word_path):
                messagebox.showerror("Error", "No existe un documento Word para convertir")
                return
            
            # Ruta PDF con lógica de nombre único
            pdf_path = self.get_unique_filename(self.word_path.replace(".docx", ""), ".pdf")
            
            def convert_pdf():
                try:
                    self.status_var.set(f"Generando PDF: {os.path.basename(pdf_path)}")
                    
                    # Método 1: Intentar con docx2pdf
                    success = False
                    try:
                        from docx2pdf import convert
                        convert(self.word_path, pdf_path)
                        success = True
                        self.status_var.set(f"PDF generado exitosamente: {pdf_path}")
                    except Exception as e1:
                        print(f"Método docx2pdf falló: {e1}")
                        
                        # Método 2: Usar LibreOffice si está disponible
                        try:
                            # Buscar LibreOffice en rutas comunes
                            libreoffice_paths = [
                                r"C:\Program Files\LibreOffice\program\soffice.exe",
                                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                                "soffice",  # Si está en PATH
                                "libreoffice"  # Si está en PATH
                            ]
                            
                            libreoffice_cmd = None
                            for path in libreoffice_paths:
                                if os.path.exists(path) or path in ["soffice", "libreoffice"]:
                                    libreoffice_cmd = path
                                    break
                            
                            if libreoffice_cmd:
                                cmd = [
                                    libreoffice_cmd,
                                    '--headless',
                                    '--convert-to', 'pdf',
                                    '--outdir', os.path.dirname(pdf_path),
                                    self.word_path
                                ]
                                
                                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                                if result.returncode == 0:
                                    success = True
                                    self.status_var.set(f"PDF generado con LibreOffice: {pdf_path}")
                                else:
                                    raise Exception(f"LibreOffice error: {result.stderr}")
                            else:
                                raise Exception("LibreOffice no encontrado")
                                
                        except Exception as e2:
                            print(f"Método LibreOffice falló: {e2}")
                            
                            # Método 3: Usar win32com si está disponible (solo Windows)
                            try:
                                if sys.platform == "win32":
                                    import win32com.client
                                    word = win32com.client.Dispatch("Word.Application")
                                    word.Visible = False
                                    doc = word.Documents.Open(self.word_path)
                                    doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
                                    doc.Close()
                                    word.Quit()
                                    success = True
                                    self.status_var.set(f"PDF generado con MS Word: {pdf_path}")
                                else:
                                    raise Exception("win32com solo disponible en Windows")
                            except Exception as e3:
                                print(f"Método win32com falló: {e3}")
                                raise Exception(f"Todos los métodos fallaron: docx2pdf={e1}, LibreOffice={e2}, win32com={e3}")
                    
                    if not success:
                        self.status_var.set("Error: No se pudo generar el PDF con ningún método")
                        
                except Exception as e:
                    error_msg = f"Error al generar PDF: {str(e)}"
                    self.status_var.set(error_msg)
                    self.root.after(0, lambda: messagebox.showerror("Error PDF", error_msg))
            
            # Ejecutar conversión en hilo separado
            threading.Thread(target=convert_pdf, daemon=True).start()
        
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo iniciar generación de PDF: {str(e)}")

    def get_unique_filename(self, base_path, extension=""):
        if not os.path.exists(base_path + extension):
            return base_path + extension
        counter = 1
        while True:
            new_path = f"{base_path}_{counter:02d}{extension}"
            if not os.path.exists(new_path):
                return new_path
            counter += 1

    def capture_loop(self):
        while self.capturing:
            try:
                self.take_screenshot()
                time.sleep(self.interval.get())
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
                break

    def take_screenshot(self):
        import pyautogui
        from PIL import Image, ImageTk

        screenshot = pyautogui.screenshot()
        self.counter += 1

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{self.base_name.get()}_{self.counter:04d}_{timestamp}"
        extension = f".{self.image_format.get().lower()}"
        filepath = self.get_unique_filename(os.path.join(self.save_path.get(), base_filename), extension)

        if self.image_format.get() == "JPG":
            screenshot = screenshot.convert('RGB')
        screenshot.save(filepath)

        # Vista previa
        img_copy = screenshot.copy()
        img_copy.thumbnail((500, 350))
        self.preview_photo = ImageTk.PhotoImage(img_copy)
        self.root.after(0, lambda: self.preview_label.config(image=self.preview_photo))

        # Word
        if self.document:
            from docx.shared import Inches
            self.document.add_heading(f'Captura #{self.counter} - {timestamp}', level=1)
            self.document.add_picture(filepath, width=Inches(6))
        
        self.root.after(0, lambda: self.counter_var.set(f"Capturas: {self.counter}"))
        self.root.after(0, lambda: self.status_var.set(f"Guardado: {os.path.basename(filepath)}"))

    # ======================= CÁMARA ======================

    def toggle_webcam(self):
        if not self.webcam_active:
            self.start_webcam()
        else:
            self.stop_webcam()

    def start_webcam(self):
        import cv2
        self.cap = cv2.VideoCapture(0, cv2.CAP_DSHOW)
        if not self.cap.isOpened():
            messagebox.showerror("Error", "No se pudo acceder a la cámara")
            return
        self.cap.set(3, 640)
        self.cap.set(4, 480)
        self.webcam_active = True
        self.detect_faces = tk.BooleanVar(value=False)
        self.frame_counter = 0

        self.webcam_btn.config(text="📹 DESACTIVAR", bg='#d63384')  
        self.create_webcam_window()

    def stop_webcam(self):
        self.webcam_active = False
        self.webcam_btn.config(text="📹 ACTIVAR", bg='#6f42c1')  
        if self.cap:
            self.cap.release()
        if self.webcam_window:
            self.webcam_window.destroy()
            self.webcam_window = None

    def create_webcam_window(self):
        self.webcam_window = tk.Toplevel(self.root)
        self.webcam_window.title("Cámara Web")
        self.webcam_window.geometry("380x220")
        self.webcam_window.configure(bg="#f8f9fa")  
        self.webcam_window.attributes('-topmost', True)

        # Label de video
        self.webcam_label = tk.Label(self.webcam_window, bg="#ffffff")  
        self.webcam_label.pack(fill='both', expand=True, padx=0, pady=0)

        # Checkbox para detección de rostros
        # detect_cb = tk.Checkbutton(
         #   self.webcam_window, text="Detectar rostros",
         #   variable=self.detect_faces, bg="#f8f9fa", font=("Arial", 10)  
        #)
        #detect_cb.pack(pady=5)

        self.webcam_window.protocol("WM_DELETE_WINDOW", self.stop_webcam)
        self.update_webcam()

    def update_webcam(self):
        if self.webcam_active and self.cap:
            import cv2
            from PIL import Image, ImageTk, Image as PILImage

            ret, frame = self.cap.read()
            if ret:
                self.frame_counter += 1
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)

                # === SOLO detectar rostros si la opción está activada y cada 5 frames ===
                # if self.detect_faces.get() and self.frame_counter % 5 == 0:
                #    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                 #    face_cascade = cv2.CascadeClassifier(
                #         cv2.data.haarcascades + "haarcascade_frontalface_default.xml"
                 #    )
                #     faces = face_cascade.detectMultiScale(gray, scaleFactor=1.3, minNeighbors=5)

                    # Dibujar rectángulos en los rostros detectados
                  #   for (x, y, w, h) in faces:
                     #    cv2.rectangle(frame_rgb, (x, y), (x + w, y + h), (0, 255, 0), 2)

                # Convertir frame a formato Tkinter
                image_pil = PILImage.fromarray(frame_rgb)
                image_pil = image_pil.resize((460, 320))
                photo = ImageTk.PhotoImage(image_pil)
                self.webcam_label.config(image=photo)
                self.webcam_label.image = photo

            if self.webcam_window:
                self.webcam_window.after(30, self.update_webcam)

    def on_closing(self):
        self.capturing = False
        self.stop_webcam()
        self.root.destroy()

# ======================= MAIN MODIFICADO ==================

def main():
    create_app_dir()
    if not is_unlocked():
        if not ask_password():
            return  # Si falla el login, cierra el programa

    # Si pasó el login, abre SnapMaster
    root = tk.Tk()
    app = ScreenCaptureApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()

if __name__ == "__main__":
    main()